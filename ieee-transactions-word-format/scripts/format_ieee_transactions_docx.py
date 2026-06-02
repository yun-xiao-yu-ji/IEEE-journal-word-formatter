#!/usr/bin/env python
"""Create an IEEE Transactions-style Word document from a source manuscript."""

from __future__ import annotations

import argparse
import re
import shutil
import tempfile
import zipfile
from copy import deepcopy
from io import BytesIO
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from lxml import etree


SKILL_DIR = Path(__file__).resolve().parents[1]
TEMPLATE_PATH = SKILL_DIR / "assets" / "ieee-transactions-template.docx"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}

MAIN_SECTION_NAMES = {
    "INTRODUCTION",
    "BACKGROUND",
    "RELATED WORK",
    "STUDY AREA AND DATA",
    "METHOD",
    "METHODS",
    "METHODOLOGY",
    "RESULT",
    "RESULTS",
    "DISCUSSION",
    "CONCLUSION",
    "CONCLUSIONS",
    "APPENDIX",
    "ACKNOWLEDGMENT",
    "ACKNOWLEDGMENTS",
    "REFERENCES",
}


def clean(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("\u00a0", " ")).strip()


def strip_numbering(text: str) -> str:
    text = clean(text)
    text = re.sub(r"^[IVXLCDM]+\.\s+", "", text, flags=re.I)
    text = re.sub(r"^\d+(\.\d+)*\.?\s+", "", text)
    text = re.sub(r"^[A-Z]\.\s+", "", text)
    return clean(text)


def roman(num: int) -> str:
    values = [(1000, "M"), (900, "CM"), (500, "D"), (400, "CD"), (100, "C"), (90, "XC"), (50, "L"), (40, "XL"), (10, "X"), (9, "IX"), (5, "V"), (4, "IV"), (1, "I")]
    out = []
    for value, symbol in values:
        while num >= value:
            out.append(symbol)
            num -= value
    return "".join(out)


def clear_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def ensure_cols(section, *, two_column: bool) -> None:
    cols = section._sectPr.xpath("./w:cols")
    if cols:
        cols = cols[0]
    else:
        cols = OxmlElement("w:cols")
        section._sectPr.append(cols)
    if two_column:
        cols.set(qn("w:num"), "2")
        cols.set(qn("w:space"), "720")
        cols.set(qn("w:equalWidth"), "0")
        for child in list(cols):
            cols.remove(child)
        col1 = OxmlElement("w:col")
        col1.set(qn("w:w"), "5040")
        col1.set(qn("w:space"), "288")
        col2 = OxmlElement("w:col")
        col2.set(qn("w:w"), "5040")
        col2.set(qn("w:space"), "0")
        cols.append(col1)
        cols.append(col2)
    else:
        for attr in (qn("w:num"), qn("w:equalWidth")):
            cols.attrib.pop(attr, None)
        for child in list(cols):
            cols.remove(child)
        cols.set(qn("w:space"), "720")


def media_blobs(docx_path: Path) -> list[bytes]:
    with zipfile.ZipFile(docx_path) as zf:
        names = sorted(
            [
                n
                for n in zf.namelist()
                if n.startswith("word/media/") and n.lower().endswith((".jpeg", ".jpg", ".png"))
            ],
            key=lambda n: [int(x) if x.isdigit() else x for x in re.split(r"(\d+)", n)],
        )
        return [zf.read(n) for n in names]


def add_paragraph(doc: Document, text: str = "", style: str | None = None):
    para = doc.add_paragraph(style=style if style and style in doc.styles else None)
    if text:
        para.add_run(text)
    return para


def add_footnote_marker(doc: Document) -> None:
    para = add_paragraph(doc, style="Text")
    para.paragraph_format.first_line_indent = 0
    run = para.add_run()
    ref = OxmlElement("w:footnoteReference")
    ref.set(qn("w:customMarkFollows"), "1")
    ref.set(qn("w:id"), "1")
    run._r.append(ref)
    sym = OxmlElement("w:sym")
    sym.set(qn("w:font"), "Symbol")
    sym.set(qn("w:char"), "F020")
    run._r.append(sym)


def add_image(doc: Document, images: list[bytes], idx: int) -> int:
    if idx >= len(images):
        return idx
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(BytesIO(images[idx]), width=Inches(3.45))
    return idx + 1


def normalize_author_line(text: str) -> str:
    text = re.sub(r"(?<=[A-Za-z])(?:[a-f](?:,[a-f])*)(?=,|\*|\s|$)", "", text)
    text = re.sub(r"\s+,", ",", text)
    text = re.sub(r",\s*,", ",", text)
    text = clean(text)
    if "," in text and " and " not in text:
        parts = [p.strip() for p in text.split(",") if p.strip()]
        if len(parts) > 1:
            text = ", ".join(parts[:-1]) + ", and " + parts[-1]
    return text


def is_affiliation(text: str) -> bool:
    return bool(re.match(r"^[a-f]\s+", text)) or any(token in text.lower() for token in ["university", "laboratory", "college", "institute", "ministry"])


def is_reference_heading(text: str) -> bool:
    return strip_numbering(text).rstrip(":").upper() in {"REFERENCES", "REFERENCE"}


def is_main_heading(text: str, style_name: str) -> bool:
    core = strip_numbering(text).upper()
    return "HEADING 1" in style_name.upper() or bool(re.match(r"^[IVXLCDM]+\.\s+\S", text, flags=re.I)) or core in MAIN_SECTION_NAMES or bool(re.match(r"^[1-9]\s+[A-Za-z]", text))


def is_subheading(text: str, style_name: str) -> bool:
    return "HEADING 2" in style_name.upper() or bool(re.match(r"^([1-9]\.[1-9]|[A-Z]\.)\s+\S", text))


def heading_text(text: str, section_no: int) -> tuple[str, int]:
    core = strip_numbering(text)
    if core.upper() in {"ACKNOWLEDGMENT", "ACKNOWLEDGMENTS", "APPENDIX"}:
        return core.title(), section_no
    return f"{roman(section_no)}. {core.upper()}", section_no + 1


def subheading_text(text: str, sub_no: int) -> str:
    core = strip_numbering(text)
    return f"{chr(ord('A') + sub_no - 1)}. {core}"


def normalize_reference(text: str, ref_no: int) -> str:
    text = clean(text)
    match = re.match(r"^\[(\d+)\]\s*(.+)", text)
    if match:
        return f"[{match.group(1)}] {match.group(2)}"
    match = re.match(r"^(\d+)\.\s*(.+)", text)
    if match:
        return f"[{match.group(1)}] {match.group(2)}"
    return f"[{ref_no}] {text}"


def add_bookmark(para, ref_no: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(1000 + ref_no))
    start.set(qn("w:name"), f"ref_{ref_no}")
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(1000 + ref_no))
    para._p.insert(0, start)
    para._p.append(end)


def add_text_with_links(para, text: str, max_ref: int) -> None:
    pattern = re.compile(r"\[(\d+(?:\s*[-,]\s*\d+)*)\]")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            para.add_run(text[pos : match.start()])
        nums = re.findall(r"\d+", match.group(1))
        if nums and all(1 <= int(n) <= max_ref for n in nums):
            link = OxmlElement("w:hyperlink")
            link.set(qn("w:anchor"), f"ref_{nums[0]}")
            run = OxmlElement("w:r")
            text_el = OxmlElement("w:t")
            text_el.text = match.group(0)
            run.append(text_el)
            link.append(run)
            para._p.append(link)
        else:
            para.add_run(match.group(0))
        pos = match.end()
    if pos < len(text):
        para.add_run(text[pos:])


def add_linked_paragraph(doc: Document, text: str, style: str, max_ref: int):
    para = add_paragraph(doc, style=style)
    add_text_with_links(para, text, max_ref)
    return para


def parse_manuscript(source: Document, args):
    items = [{"text": clean(p.text), "style": p.style.name} for p in source.paragraphs if clean(p.text)]
    title = args.title or (items[0]["text"] if items else "")
    authors = args.authors or (items[1]["text"] if len(items) > 1 else "")
    affiliations: list[str] = []
    abstract = args.abstract or ""
    keywords = args.keywords or ""
    funding = ""
    body_start = 2
    references_at = None

    for idx, item in enumerate(items[2:25], start=2):
        text = item["text"]
        lower = text.lower()
        if lower.startswith("abstract"):
            abstract = abstract or re.sub(r"^abstract\s*[:\-\u2014]*\s*", "", text, flags=re.I)
            body_start = idx + 1
        elif lower.startswith(("keyword", "index terms")):
            keywords = keywords or re.sub(r"^(keywords?|index terms)\s*[:\-\u2014]*\s*", "", text, flags=re.I)
            body_start = idx + 1
            break
        elif is_affiliation(text):
            affiliations.append(text)

    for idx, item in enumerate(items):
        text = item["text"]
        lower = text.lower()
        if lower == "funding" and idx + 1 < len(items):
            funding = items[idx + 1]["text"].rstrip(".") + "."
        if is_reference_heading(text):
            references_at = idx
            break

    body_end = references_at if references_at is not None else len(items)
    references = [i["text"] for i in items[(references_at + 1) if references_at is not None else len(items) :]]
    return title, normalize_author_line(authors), affiliations, abstract, keywords, funding, items[body_start:body_end], references


def patch_docx(output_path: Path, footnote_paragraphs: list[str]) -> None:
    tmp = Path(tempfile.mkdtemp())
    try:
        unpacked = tmp / "docx"
        with zipfile.ZipFile(output_path) as zin:
            zin.extractall(unpacked)
        styles_path = unpacked / "word" / "styles.xml"
        document_path = unpacked / "word" / "document.xml"
        footnotes_path = unpacked / "word" / "footnotes.xml"

        styles = styles_path.read_text(encoding="utf-8")
        for style_id in ("Title", "Authors"):
            styles = re.sub(rf'(<w:style\b[^>]*w:styleId="{style_id}"[^>]*>.*?</w:style>)', lambda m: re.sub(r"\s*<w:framePr\b[^>]*/>", "", m.group(1)), styles, flags=re.S)
        styles = re.sub(r'(<w:style\b[^>]*w:styleId="References"[^>]*>.*?</w:style>)', lambda m: re.sub(r"\s*<w:numPr>.*?</w:numPr>", "", m.group(1), flags=re.S), styles, flags=re.S)
        styles_path.write_text(styles, encoding="utf-8")

        tree = etree.parse(str(document_path))
        body = tree.getroot().find(".//w:body", namespaces=NS)
        paragraphs = body.findall("w:p", namespaces=NS)
        footnote_para = next((p for p in paragraphs[:12] if p.find(".//w:footnoteReference", namespaces=NS) is not None), None)
        if footnote_para is not None:
            ppr = footnote_para.find("w:pPr", namespaces=NS)
            sect_pr = ppr.find("w:sectPr", namespaces=NS) if ppr is not None else None
            if sect_pr is not None:
                sect_copy = deepcopy(sect_pr)
                ppr.remove(sect_pr)
                break_para = etree.Element(f"{{{W_NS}}}p")
                break_ppr = etree.SubElement(break_para, f"{{{W_NS}}}pPr")
                break_ppr.append(sect_copy)
                body.insert(body.index(footnote_para), break_para)
        tree.write(str(document_path), encoding="UTF-8", xml_declaration=True, standalone=True)

        footnotes = footnotes_path.read_text(encoding="utf-8")
        paras = []
        for text in footnote_paragraphs:
            paras.append(
                '<w:p><w:pPr><w:pStyle w:val="footnote text"/><w:keepLines/><w:ind w:firstLine="202"/><w:jc w:val="both"/>'
                '<w:rPr><w:sz w:val="16"/><w:szCs w:val="16"/></w:rPr></w:pPr><w:r><w:rPr><w:sz w:val="16"/><w:szCs w:val="16"/></w:rPr>'
                f'<w:t xml:space="preserve">{escape(text)}</w:t></w:r></w:p>'
            )
        footnotes = re.sub(r'<w:footnote w:id="1">.*?</w:footnote>', '<w:footnote w:id="1">' + "".join(paras) + "</w:footnote>", footnotes, flags=re.S)
        footnotes_path.write_text(footnotes, encoding="utf-8")

        with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for path in unpacked.rglob("*"):
                if path.is_file():
                    zout.write(path, path.relative_to(unpacked).as_posix())
    finally:
        shutil.rmtree(tmp)


def build_document(input_path: Path, output_path: Path, args) -> list[str]:
    source = Document(input_path)
    output = Document(TEMPLATE_PATH)
    clear_body(output)
    ensure_cols(output.sections[0], two_column=False)
    images = media_blobs(input_path)
    warnings = []

    title, authors, affiliations, abstract, keywords, funding, body_items, references = parse_manuscript(source, args)
    max_ref = len(references)

    add_paragraph(output, title, "Title")
    add_paragraph(output, authors, "Authors")
    output.add_section(WD_SECTION.CONTINUOUS)
    ensure_cols(output.sections[0], two_column=False)
    ensure_cols(output.sections[-1], two_column=True)
    add_footnote_marker(output)
    if abstract:
        add_paragraph(output, f"Abstract\u2014{abstract}", "Abstract")
    if keywords:
        add_paragraph(output, f"Index Terms\u2014{keywords}", "IndexTerms")

    section_no = 1
    sub_no = 0
    image_idx = 0
    first_intro_done = False

    for item in body_items:
        text = item["text"]
        style = item["style"]
        lower = text.lower()
        if lower.startswith(("abstract", "keyword", "index terms")) or lower == "funding":
            continue
        if re.match(r"^(figure|fig\.)\s*\d+", text, flags=re.I):
            image_idx = add_image(output, images, image_idx)
            add_paragraph(output, re.sub(r"^Figure\s+", "Fig. ", text, count=1, flags=re.I), "Figure Caption")
        elif is_main_heading(text, style):
            heading, section_no = heading_text(text, section_no)
            add_paragraph(output, heading, "Heading 1")
            sub_no = 0
        elif is_subheading(text, style):
            sub_no += 1
            add_paragraph(output, subheading_text(text, sub_no), "Heading 2")
        elif output.paragraphs and output.paragraphs[-1].text == "I. INTRODUCTION" and not first_intro_done:
            add_paragraph(output, text[:1].upper(), "Normal")
            add_linked_paragraph(output, text[1:], "Text", max_ref)
            first_intro_done = True
        else:
            add_linked_paragraph(output, text, "Text", max_ref)

    if references:
        add_paragraph(output, "References", "Reference Head")
        for idx, ref in enumerate(references, start=1):
            para = add_paragraph(output, normalize_reference(ref, idx), "References")
            add_bookmark(para, idx)

    if source.tables:
        warnings.append(f"Source contains {len(source.tables)} table(s); inspect tables manually.")
    if len(source.inline_shapes) != image_idx:
        warnings.append(f"Source contains {len(source.inline_shapes)} inline object(s); inserted {image_idx} supported image(s).")

    output.save(output_path)
    footnote_paragraphs = [funding or "Manuscript received [date]; revised [date]; accepted [date]."]
    footnote_paragraphs.append("(Corresponding author information should be confirmed before submission.)")
    footnote_paragraphs.extend(affiliations)
    patch_docx(output_path, footnote_paragraphs)
    return warnings


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input_docx", type=Path)
    parser.add_argument("output_docx", type=Path)
    parser.add_argument("--title")
    parser.add_argument("--authors")
    parser.add_argument("--abstract")
    parser.add_argument("--keywords")
    args = parser.parse_args()

    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Missing template: {TEMPLATE_PATH}")
    if not args.input_docx.exists():
        raise FileNotFoundError(f"Missing input document: {args.input_docx}")

    args.output_docx.parent.mkdir(parents=True, exist_ok=True)
    warnings = build_document(args.input_docx, args.output_docx, args)
    print(f"Wrote: {args.output_docx}")
    for warning in warnings:
        print(f"WARNING: {warning}")


if __name__ == "__main__":
    main()
